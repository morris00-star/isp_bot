from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


class DjangoChatbotGuideGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document_styles()

    def setup_document_styles(self):
        """Setup professional document styles"""
        # Normal text style
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Heading 1 style (Chapters)
        style = self.doc.styles['Heading 1']
        style.font.name = 'Calibri'
        style.font.size = Pt(18)
        style.font.bold = True
        style.font.color.rgb = RGBColor(44, 62, 80)

        # Heading 2 style (Files)
        style = self.doc.styles['Heading 2']
        style.font.name = 'Calibri'
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = RGBColor(52, 73, 94)

        # Heading 3 style (Subsections)
        style = self.doc.styles['Heading 3']
        style.font.name = 'Calibri'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(52, 152, 219)

        # Code block style
        style = self.doc.styles.add_style('CodeBlock', 1)
        style.font.name = 'Courier New'
        style.font.size = Pt(10)

    def add_heading_centered(self, text, level=1):
        """Add centered heading"""
        heading = self.doc.add_heading(text, level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return heading

    def add_horizontal_line(self):
        """Add horizontal line"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.add_run('_' * 70)

    def add_code_block(self, code, language="python"):
        """Add formatted code block"""
        p = self.doc.add_paragraph()
        p.style = 'CodeBlock'
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(8)

        # Add background shading (light gray)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "F5F5F5")
        p._element.get_or_add_pPr().append(shading_elm)

        # Split code and add each line
        for line in code.strip().split('\n'):
            run = p.add_run(line + '\n')
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    def add_bullet_point(self, text, bold=False):
        """Add bullet point"""
        p = self.doc.add_paragraph(text, style='List Bullet')
        if bold:
            p.add_run().bold = True
        return p

    def create_title_page(self):
        """Create professional title page"""
        # Add empty paragraphs for spacing
        for _ in range(6):
            self.doc.add_paragraph()

        # Main title
        title = self.add_heading_centered('DJANGO ROUTER SUPPORT CHATBOT', 0)
        title.runs[0].font.size = Pt(32)
        title.runs[0].font.color.rgb = RGBColor(102, 126, 234)

        self.doc.add_paragraph()

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Complete Implementation Manual')
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(118, 75, 162)

        for _ in range(4):
            self.doc.add_paragraph()

        # Author information
        author_info = [
            ('Author:', 'Senior Django Developer'),
            ('Phone:', '+1 (555) 123-4567'),
            ('Email:', 'django.developer@example.com'),
            ('Address:', '123 Tech Street, Silicon Valley, CA 94025'),
            ('Date:', datetime.now().strftime('%B %d, %Y'))
        ]

        for label, value in author_info:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'{label} ')
            run.bold = True
            run.font.size = Pt(12)
            run = p.add_run(value)
            run.font.size = Pt(12)

        self.doc.add_page_break()

    def create_table_of_contents(self):
        """Create table of contents"""
        self.add_heading_centered('TABLE OF CONTENTS', 1)

        toc_items = [
            'CHAPTER 1: PROJECT OVERVIEW ................................ 4',
            'CHAPTER 2: PROJECT SETUP ...................................... 5',
            'CHAPTER 3: DJANGO CONFIGURATION FILES ................ 8',
            'CHAPTER 4: CHATBOT APPLICATION FILES .................. 15',
            'CHAPTER 5: TEMPLATES AND FRONTEND ........................ 42',
            'CHAPTER 6: STATIC FILES AND STYLING ........................ 58',
            'CHAPTER 7: DEPLOYMENT CONFIGURATION .................... 72',
            'CHAPTER 8: FEATURES SUMMARY .................................... 85'
        ]

        for item in toc_items:
            p = self.doc.add_paragraph()
            p.add_run(item).font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(6)

        self.doc.add_page_break()

    def chapter1_overview(self):
        """Chapter 1: Project Overview"""
        self.add_heading_centered('CHAPTER 1', 1)
        self.add_heading_centered('PROJECT OVERVIEW', 2)

        # What the project does
        self.doc.add_heading('1.1 What the Project Does', 3)
        self.doc.add_paragraph(
            'The Router Support Chatbot is a Django-based web application designed to provide '
            'automated customer support for Huawei A5 and V5 router users. It features an intelligent '
            'rule-based chatbot that can answer common questions about router setup, troubleshooting, '
            'and configuration.'
        )

        # Problem it solves
        self.doc.add_heading('1.2 Problem It Solves', 3)
        self.doc.add_paragraph(
            'This application addresses the following challenges:'
        )
        problems = [
            'Reduces customer support workload by automating responses to common router queries',
            'Provides 24/7 instant support availability without human intervention',
            'Offers consistent and accurate information across all user interactions',
            'Helps users troubleshoot router issues independently',
            'Reduces response time from minutes/hours to milliseconds',
            'Maintains conversation history for better user experience'
        ]
        for problem in problems:
            self.add_bullet_point(problem)

        # Technologies used
        self.doc.add_heading('1.3 Technologies Used', 3)
        tech_stack = [
            ('Backend Framework:', 'Django 4.2+ - Python web framework'),
            ('Frontend:', 'HTML5, CSS3, JavaScript (ES6+)'),
            ('Database:', 'SQLite (development), PostgreSQL (production)'),
            ('Server:', 'Gunicorn WSGI server'),
            ('Deployment:', 'Render.com cloud platform'),
            ('Monitoring:', 'UptimeRobot for 24/7 uptime monitoring'),
            ('Static Files:', 'Whitenoise for production static file serving'),
            ('Version Control:', 'Git with GitHub repository'
             )]

        for tech, desc in tech_stack:
            p = self.doc.add_paragraph()
            p.add_run(f'{tech} ').bold = True
            p.add_run(desc)

        # System architecture overview
        self.doc.add_heading('1.4 System Architecture Overview', 3)
        architecture = '''
The application follows Django's Model-View-Template (MVT) architecture:

1. Client Layer:
   - Web browser (mobile/desktop)
   - Makes HTTP requests to the server
   - Renders HTML/CSS/JavaScript responses

2. Server Layer (Django):
   - URL dispatcher routes requests to appropriate views
   - Views process requests and implement business logic
   - Templates generate dynamic HTML responses
   - Session management stores conversation history

3. Data Layer:
   - JSON file stores chatbot responses and intents
   - Session storage maintains chat history
   - Database stores application data (optional)

4. Chatbot Logic:
   - Rule-based intent matching system
   - Keyword extraction and matching algorithm
   - Fallback responses for unmatched queries
   - Context-aware response generation

5. Deployment Infrastructure:
   - Gunicorn WSGI server processes requests
   - Whitenoise serves static files efficiently
   - PostgreSQL for production database
   - UptimeRobot keeps application awake
'''
        self.doc.add_paragraph(architecture)
        self.doc.add_page_break()

    def chapter2_setup(self):
        """Chapter 2: Project Setup"""
        self.add_heading_centered('CHAPTER 2', 1)
        self.add_heading_centered('PROJECT SETUP AND INSTALLATION', 2)

        # Environment setup
        self.doc.add_heading('2.1 Environment Setup', 3)
        self.doc.add_paragraph(
            'Before starting the Django project, you need to set up a proper development environment. '
            'This ensures project dependencies are isolated and manageable.'
        )

        self.doc.add_heading('2.1.1 Virtual Environment Creation', 4)
        setup_steps = [
            ('Windows (Command Prompt):', 'python -m venv my_venv'),
            ('Windows (PowerShell):', 'python -m venv my_venv'),
            ('macOS/Linux:', 'python3 -m venv my_venv')
        ]

        for os_name, command in setup_steps:
            p = self.doc.add_paragraph()
            p.add_run(f'{os_name} ').bold = True
            self.add_code_block(command, "bash")

        self.doc.add_heading('2.1.2 Virtual Environment Activation', 4)
        activation_steps = [
            ('Windows (Command Prompt):', 'my_venv\\Scripts\\activate.bat'),
            ('Windows (PowerShell):', 'my_venv\\Scripts\\Activate.ps1'),
            ('macOS/Linux:', 'source my_venv/bin/activate')
        ]

        for os_name, command in activation_steps:
            p = self.doc.add_paragraph()
            p.add_run(f'{os_name} ').bold = True
            self.add_code_block(command, "bash")

        # Django installation
        self.doc.add_heading('2.2 Django Installation', 3)
        self.add_code_block('pip install django', "bash")

        # Project creation
        self.doc.add_heading('2.3 Project Creation', 3)
        commands = [
            'django-admin startproject router_chatbot',
            'cd router_chatbot'
        ]
        for cmd in commands:
            self.add_code_block(cmd, "bash")

        # App creation
        self.doc.add_heading('2.4 App Creation', 3)
        self.add_code_block('python manage.py startapp chatbot', "bash")

        self.doc.add_heading('2.5 Project Structure Explanation', 3)
        structure = '''
router_chatbot/                  # Project root directory
├── manage.py                    # Django's command-line utility
├── router_chatbot/              # Project configuration directory
│   ├── __init__.py             # Python package marker
│   ├── settings.py              # Project settings
│   ├── urls.py                  # Main URL configuration
│   └── wsgi.py                   # WSGI configuration for deployment
└── chatbot/                      # Main application directory
    ├── __init__.py               # Python package marker
    ├── admin.py                   # Admin interface configuration
    ├── apps.py                    # App configuration
    ├── models.py                  # Database models
    ├── views.py                   # View functions
    ├── urls.py                    # App URL configuration
    ├── data/                       # JSON data files
    │   └── responses.json          # Chatbot responses
    ├── templates/                  # HTML templates
    │   └── chatbot/
    │       └── chat.html           # Main chat interface
    └── static/                      # Static files
        └── chatbot/
            ├── style.css             # CSS styles
            └── manifest.json          # PWA manifest
'''
        self.doc.add_paragraph(structure)
        self.doc.add_page_break()

    def chapter3_configuration(self):
        """Chapter 3: Django Configuration Files"""
        self.add_heading_centered('CHAPTER 3', 1)
        self.add_heading_centered('DJANGO CONFIGURATION FILES', 2)

        # settings.py
        self.doc.add_heading('FILE: router_chatbot/settings.py', 2)

        self.doc.add_heading('3.1 Purpose of settings.py', 3)
        self.doc.add_paragraph(
            'settings.py is the central configuration file for the entire Django project. '
            'It controls database connections, installed applications, middleware, template engines, '
            'static files, security settings, and all other project-wide configurations.'
        )

        self.doc.add_heading('3.2 Why This File is Needed', 3)
        reasons = [
            'Defines the environment (development/production) through DEBUG flag',
            'Lists all installed applications for Django to recognize',
            'Configures middleware for request/response processing',
            'Sets up database connections and authentication',
            'Manages static and media file handling',
            'Controls security settings like allowed hosts and CORS',
            'Defines template locations and context processors'
        ]
        for reason in reasons:
            self.add_bullet_point(reason)

        self.doc.add_heading('3.3 Full settings.py Code', 3)
        settings_code = '''import os
from pathlib import Path
import dj_database_url
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Build paths inside the project
BASE_DIR = Path(__file__).resolve().parent.parent

# Security settings
SECRET_KEY = os.getenv('SECRET_KEY', 'django-insecure-your-dev-key-here')
DEBUG = os.getenv('DEBUG', 'False') == 'True'
ALLOWED_HOSTS = ['localhost', '127.0.0.1', '.onrender.com']

# Application definition
INSTALLED_APPS = [
    'django.contrib.admin',
    'django.contrib.auth',
    'django.contrib.contenttypes',
    'django.contrib.sessions',
    'django.contrib.messages',
    'whitenoise.runserver_nostatic',
    'django.contrib.staticfiles',
    'chatbot',  # Our chatbot application
]

MIDDLEWARE = [
    'django.middleware.security.SecurityMiddleware',
    'whitenoise.middleware.WhiteNoiseMiddleware',
    'django.contrib.sessions.middleware.SessionMiddleware',
    'django.middleware.common.CommonMiddleware',
    'django.middleware.csrf.CsrfViewMiddleware',
    'django.contrib.auth.middleware.AuthenticationMiddleware',
    'django.contrib.messages.middleware.MessageMiddleware',
    'django.middleware.clickjacking.XFrameOptionsMiddleware',
]

ROOT_URLCONF = 'router_chatbot.urls'

TEMPLATES = [
    {
        'BACKEND': 'django.template.backends.django.DjangoTemplates',
        'DIRS': [],
        'APP_DIRS': True,
        'OPTIONS': {
            'context_processors': [
                'django.template.context_processors.debug',
                'django.template.context_processors.request',
                'django.contrib.auth.context_processors.auth',
                'django.contrib.messages.context_processors.messages',
            ],
        },
    },
]

WSGI_APPLICATION = 'router_chatbot.wsgi.application'

# Database configuration
if DEBUG:
    DATABASES = {
        'default': {
            'ENGINE': 'django.db.backends.sqlite3',
            'NAME': BASE_DIR / 'db.sqlite3',
        }
    }
else:
    DATABASES = {
        'default': dj_database_url.config(
            default=os.getenv('DATABASE_URL'),
            conn_max_age=600
        )
    }

# Password validation
AUTH_PASSWORD_VALIDATORS = [
    {
        'NAME': 'django.contrib.auth.password_validation.UserAttributeSimilarityValidator',
    },
    {
        'NAME': 'django.contrib.auth.password_validation.MinimumLengthValidator',
    },
    {
        'NAME': 'django.contrib.auth.password_validation.CommonPasswordValidator',
    },
    {
        'NAME': 'django.contrib.auth.password_validation.NumericPasswordValidator',
    },
]

# Internationalization
LANGUAGE_CODE = 'en-us'
TIME_ZONE = 'UTC'
USE_I18N = True
USE_TZ = True

# Static files (CSS, JavaScript, Images)
STATIC_URL = '/static/'
STATIC_ROOT = BASE_DIR / 'staticfiles'
STATICFILES_DIRS = [BASE_DIR / 'static']
STATICFILES_STORAGE = 'whitenoise.storage.CompressedManifestStaticFilesStorage'

# Security settings for production
if not DEBUG:
    SECURE_SSL_REDIRECT = True
    SESSION_COOKIE_SECURE = True
    CSRF_COOKIE_SECURE = True
    SECURE_BROWSER_XSS_FILTER = True
    SECURE_CONTENT_TYPE_NOSNIFF = True
    SECURE_HSTS_SECONDS = 31536000
    SECURE_HSTS_INCLUDE_SUBDOMAINS = True
    SECURE_HSTS_PRELOAD = True
    X_FRAME_OPTIONS = 'DENY'

# Default primary key field type
DEFAULT_AUTO_FIELD = 'django.db.models.BigAutoField'

# Session configuration
SESSION_ENGINE = 'django.contrib.sessions.backends.db'
SESSION_COOKIE_AGE = 86400  # 24 hours in seconds
SESSION_SAVE_EVERY_REQUEST = True'''

        self.add_code_block(settings_code, "python")

        # Line-by-line explanation
        self.doc.add_heading('3.4 Line-by-Line Explanation', 3)

        explanations = [
            ('from pathlib import Path', 'Imports Python\'s modern path handling class for cross-platform file paths'),
            ('BASE_DIR = Path(__file__).resolve().parent.parent',
             'Gets the absolute path to the project root directory'),
            ('SECRET_KEY = os.getenv(...)', 'Retrieves the secret key from environment variables for security'),
            ('DEBUG = os.getenv(\'DEBUG\', \'False\') == \'True\'',
             'Sets debug mode from environment, defaults to False for production'),
            ('INSTALLED_APPS', 'Lists all Django apps and third-party apps this project uses'),
            ('\'chatbot\',', 'Registers our custom chatbot app with Django'),
            ('MIDDLEWARE', 'Defines request/response processing pipeline'),
            ('WhitenoiseMiddleware', 'Enables efficient static file serving in production'),
            ('TEMPLATES', 'Configures Django template engine settings'),
            ('DATABASES', 'Sets up database connection (SQLite for dev, PostgreSQL for production)'),
            ('STATICFILES_STORAGE', 'Configures Whitenoise for compressed static files'),
            ('SESSION_SAVE_EVERY_REQUEST = True', 'Ensures chat history persists by saving session on every request')
        ]

        for item, desc in explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'{item}: ').bold = True
            p.add_run(desc)

        # Main urls.py
        self.doc.add_heading('FILE: router_chatbot/urls.py', 2)
        urls_code = '''from django.contrib import admin
from django.urls import path, include

urlpatterns = [
    path('admin/', admin.site.urls),  # Admin interface
    path('', include('chatbot.urls')),  # Include chatbot app URLs
]'''
        self.add_code_block(urls_code, "python")

        self.doc.add_heading('3.5 URL Routing Explanation', 3)
        url_explanations = [
            ('path(\'admin/\', admin.site.urls)', 'Routes all admin URLs to Django\'s built-in admin interface'),
            ('path(\'\', include(\'chatbot.urls\'))', 'Routes root URL (/) to chatbot app\'s URL configuration'),
            ('include() function', 'Delegates URL matching to the specified app\'s urls.py file')
        ]

        for item, desc in url_explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'{item}: ').bold = True
            p.add_run(desc)

        self.doc.add_page_break()

    def chapter4_chatbot_app(self):
        """Chapter 4: Chatbot Application Files"""
        self.add_heading_centered('CHAPTER 4', 1)
        self.add_heading_centered('CHATBOT APPLICATION FILES', 2)

        # chatbot/urls.py
        self.doc.add_heading('FILE: chatbot/urls.py', 2)
        urls_code = '''from django.urls import path
from . import views

urlpatterns = [
    path('', views.chat_view, name='chat'),  # Main chat interface
    path('get-response/', views.get_response, name='get_response'),  # AJAX endpoint for responses
    path('contact/', views.contact_view, name='contact'),  # Contact page
    path('clear-history/', views.clear_history, name='clear_history'),  # Clear session history
    path('health/', views.health_check, name='health_check'),  # Detailed health check
    path('ping/', views.ping, name='ping'),  # Simple ping endpoint for uptime monitoring
]'''
        self.add_code_block(urls_code, "python")

        self.doc.add_heading('4.1 URL Patterns Explanation', 3)
        url_patterns = [
            ('chat_view', 'Main view that renders the chat interface and displays chat history'),
            ('get_response', 'AJAX endpoint that processes user messages and returns bot responses'),
            ('contact_view', 'Simple contact information page'),
            ('clear_history', 'Clears the session-stored chat history'),
            ('health_check', 'Detailed health check with system metrics'),
            ('ping', 'Lightweight endpoint for uptime monitoring services')
        ]

        for name, desc in url_patterns:
            p = self.doc.add_paragraph()
            p.add_run(f'{name}: ').bold = True
            p.add_run(desc)

        # chatbot/views.py
        self.doc.add_heading('FILE: chatbot/views.py', 2)
        views_code = '''from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
import json
import os
from django.conf import settings
import time

# Load responses from JSON file
def load_responses():
    """
    Load chatbot responses from the JSON data file.
    Returns: dict containing all response data
    """
    json_path = os.path.join(settings.BASE_DIR, 'chatbot', 'data', 'responses.json')
    with open(json_path, 'r', encoding='utf-8') as file:
        return json.load(file)

# Load data once when module loads for performance
DATA = load_responses()
RESPONSES = DATA['responses']
FALLBACK_MESSAGE = DATA['fallback_message']
GREETINGS = DATA['greetings']
GREETING_RESPONSE = DATA['greeting_response']
THANKS = DATA['thanks']
THANKS_RESPONSE = DATA['thanks_response']
SUGGESTED_QUESTIONS = DATA['suggested_questions']

def get_bot_response(user_message):
    """
    Enhanced rule-based function to match user message with responses.

    Algorithm:
    1. Convert message to lowercase for case-insensitive matching
    2. Check for greetings first (highest priority)
    3. Check for thanks expressions
    4. Match against intent keywords using keyword counting
    5. Return best matching response or fallback message

    Args:
        user_message (str): The user's input message

    Returns:
        str: The bot's response
    """
    # Convert to lowercase for case-insensitive matching
    user_message_lower = user_message.lower().strip()

    # Check for greetings first (priority 1)
    for greeting in GREETINGS:
        if greeting in user_message_lower:
            return GREETING_RESPONSE

    # Check for thanks (priority 2)
    for thank in THANKS:
        if thank in user_message_lower:
            return THANKS_RESPONSE

    # Check each intent's keywords and count matches
    best_match = None
    max_keywords_matched = 0

    for intent in RESPONSES:
        keywords_matched = 0
        for keyword in intent['keywords']:
            if keyword in user_message_lower:
                keywords_matched += 1

        # If we found a match with more keywords than previous best
        if keywords_matched > max_keywords_matched:
            max_keywords_matched = keywords_matched
            best_match = intent['response']

    # Return best match if found (and at least 1 keyword matched)
    if best_match and max_keywords_matched > 0:
        return best_match

    # If no match found, return fallback message
    return FALLBACK_MESSAGE

def chat_view(request):
    """
    Main view to display the chat interface.

    This view:
    1. Initializes chat history in session if it doesn't exist
    2. Adds welcome message for new sessions
    3. Passes suggested questions and history to template

    Args:
        request: HTTP request object

    Returns:
        Rendered chat.html template with context
    """
    # Initialize chat history in session if it doesn't exist
    if 'chat_history' not in request.session:
        request.session['chat_history'] = []

        # Add welcome message for new sessions
        welcome_message = "Hello! 👋 I'm your router support assistant. How can I help you with your Huawei A5 or V5 router today?"
        request.session['chat_history'].append({'type': 'bot', 'text': welcome_message})
        request.session.modified = True

    context = {
        'suggested_questions': SUGGESTED_QUESTIONS,
        'chat_history': request.session['chat_history']
    }
    return render(request, 'chat.html', context)

def get_response(request):
    """
    AJAX view to handle message sending and get bot response.

    This view:
    1. Receives POST request with user message
    2. Validates and processes the message
    3. Gets bot response using rule-based matching
    4. Updates session chat history
    5. Returns JSON response to frontend

    Args:
        request: HTTP request object (must be POST)

    Returns:
        JsonResponse with bot response or error
    """
    if request.method == 'POST':
        try:
            # Get user message from POST data
            data = json.loads(request.body)
            user_message = data.get('message', '')

            if not user_message.strip():
                return JsonResponse({
                    'response': "Please type a message!",
                    'success': True
                })

            # Get bot response
            bot_response = get_bot_response(user_message)

            # Update chat history in session
            if 'chat_history' not in request.session:
                request.session['chat_history'] = []

            # Add messages to history
            chat_history = request.session['chat_history']
            chat_history.append({'type': 'user', 'text': user_message})
            chat_history.append({'type': 'bot', 'text': bot_response})

            # Keep only last 30 messages to prevent session from getting too large
            if len(chat_history) > 30:
                chat_history = chat_history[-30:]

            request.session['chat_history'] = chat_history
            request.session.modified = True

            # Return bot response as JSON
            return JsonResponse({
                'response': bot_response,
                'success': True
            })

        except json.JSONDecodeError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid JSON data'
            }, status=400)
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)

    return JsonResponse({
        'success': False,
        'error': 'Invalid request method'
    }, status=405)

def contact_view(request):
    """
    Simple view to display contact information page.

    Args:
        request: HTTP request object

    Returns:
        Rendered contact.html template
    """
    return render(request, 'contact.html')

def clear_history(request):
    """
    View to clear chat history from session.

    Args:
        request: HTTP request object

    Returns:
        JsonResponse indicating success
    """
    if 'chat_history' in request.session:
        del request.session['chat_history']
    return JsonResponse({'success': True})

def health_check(request):
    """
    Lightweight health check endpoint for uptime monitoring.
    Returns system status and basic metrics.

    Args:
        request: HTTP request object

    Returns:
        JsonResponse with health status or simple HttpResponse
    """
    if request.method == 'GET':
        try:
            # Try to include system metrics if psutil is available
            import psutil
            memory_info = psutil.virtual_memory()
            response_data = {
                'status': 'healthy',
                'timestamp': time.time(),
                'memory_usage': f"{memory_info.percent}%",
            }
            return JsonResponse(response_data)
        except ImportError:
            # Fallback to simple response if psutil not available
            return HttpResponse("OK", status=200)

    return HttpResponse("Method not allowed", status=405)

def ping(request):
    """
    Ultra-light ping endpoint - just returns 200 OK.
    Perfect for UptimeRobot monitoring.

    Args:
        request: HTTP request object

    Returns:
        HttpResponse with "OK" status
    """
    return HttpResponse("OK", content_type="text/plain")'''

        self.add_code_block(views_code, "python")

        self.doc.add_heading('4.2 View Functions Explanation', 3)

        view_explanations = [
            ('load_responses()', 'Loads JSON data file containing all chatbot responses and configurations'),
            ('get_bot_response()', 'Core chatbot logic implementing keyword-based intent matching algorithm'),
            ('chat_view()', 'Main view that initializes session and renders the chat interface'),
            ('get_response()', 'AJAX handler that processes messages and updates conversation history'),
            ('contact_view()', 'Renders contact information page'),
            ('clear_history()', 'Utility function to clear session-stored chat history'),
            ('health_check()', 'Monitoring endpoint for system health checks'),
            ('ping()', 'Lightweight endpoint for uptime monitoring services')
        ]

        for name, desc in view_explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'{name}: ').bold = True
            p.add_run(desc)

        self.doc.add_page_break()

    def chapter5_templates(self):
        """Chapter 5: Templates and Frontend"""
        self.add_heading_centered('CHAPTER 5', 1)
        self.add_heading_centered('TEMPLATES AND FRONTEND', 2)

        # chat.html
        self.doc.add_heading('FILE: chatbot/templates/chat.html', 2)
        self.doc.add_heading('5.1 Main Chat Interface Template', 3)

        chat_html = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=yes, viewport-fit=cover">
    <meta name="theme-color" content="#667eea">
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <meta name="format-detection" content="telephone=no">
    <title>Router Support Chatbot - Huawei A5 & V5</title>
    {% load static %}
    <link rel="stylesheet" href="{% static 'chatbot/style.css' %}">
    <link rel="apple-touch-icon" href="{% static 'chatbot/icon.png' %}">
    <link rel="manifest" href="{% static 'chatbot/manifest.json' %}">
</head>
<body>
    <div class="container">
        <!-- Header Section -->
        <div class="header">
            <h1>🔧 Router Support Assistant</h1>
            <button onclick="clearHistory()" class="clear-btn" title="Clear chat history">🗑️ Clear Chat</button>
        </div>
        <p class="subtitle">Support for Huawei A5 (2.4GHz) and Huawei V5 (Dual Band) routers</p>

        <!-- Main Chat Container -->
        <div class="chat-container">
            <div class="chat-messages" id="chatMessages">
                <!-- Chat history displayed here -->
                {% for message in chat_history %}
                    <div class="message {{ message.type }}-message" id="message-{{ forloop.counter }}">
                        <div class="message-content">
                            {% if message.type == 'bot' %}
                                <span class="icon">🤖</span>
                            {% else %}
                                <span class="icon">👤</span>
                            {% endif %}
                            {{ message.text|safe|linebreaksbr }}
                        </div>
                        <div class="message-time">
                            {{ message.timestamp|default:"" }}
                        </div>
                    </div>
                {% empty %}
                    <div class="message bot-message">
                        <div class="message-content">
                            <span class="icon">🤖</span>
                            Hello! 👋 I'm your router support assistant. How can I help you with your Huawei A5 or V5 router today?
                        </div>
                    </div>
                {% endfor %}

                <!-- Typing indicator (hidden by default) -->
                <div id="typingIndicator" class="message bot-message typing-indicator" style="display: none;">
                    <div class="message-content">
                        <span class="icon">🤖</span>
                        <span class="dot"></span>
                        <span class="dot"></span>
                        <span class="dot"></span>
                    </div>
                </div>
            </div>

            <!-- Input Area -->
            <div class="input-area">
                <input type="text"
                       id="userInput"
                       placeholder="Type your question here... (e.g., slow internet, change password)"
                       onkeypress="handleKeyPress(event)"
                       autocomplete="off"
                       autofocus>
                <button onclick="sendMessage()" id="sendButton">
                    <span>Send</span>
                    <span class="shortcut">↵</span>
                </button>
            </div>

            <!-- Quick Action Buttons -->
            <div class="quick-actions">
                <button onclick="useSuggestedQuestion('Why is my internet slow?')" class="quick-btn" title="Common issue">🐢 Slow Internet</button>
                <button onclick="useSuggestedQuestion('How do I change my WiFi password?')" class="quick-btn" title="Security">🔑 Change Password</button>
                <button onclick="useSuggestedQuestion('What\\'s the difference between Huawei A5 and V5?')" class="quick-btn" title="Router models">📱 Compare Models</button>
            </div>
        </div>

        <!-- Suggested Questions Section -->
        <div class="suggested-questions">
            <h3>💡 Frequently Asked Questions:</h3>
            <div class="question-categories">
                <div class="category">
                    <h4>🔌 Connection Issues</h4>
                    <div class="question-buttons">
                        <button onclick="useSuggestedQuestion('Why is my router showing only one light?')">
                            ⚡ Only one light
                        </button>
                        <button onclick="useSuggestedQuestion('Why does my internet keep disconnecting?')">
                            📶 Internet keeps dropping
                        </button>
                        <button onclick="useSuggestedQuestion('Why are my devices not connecting to WiFi?')">
                            📱 Devices won't connect
                        </button>
                    </div>
                </div>

                <div class="category">
                    <h4>⚡ Performance</h4>
                    <div class="question-buttons">
                        <button onclick="useSuggestedQuestion('Why is my internet slow?')">
                            🐢 Slow internet
                        </button>
                        <button onclick="useSuggestedQuestion('Why is WiFi weak in some rooms?')">
                            📡 Weak signal in some areas
                        </button>
                        <button onclick="useSuggestedQuestion('What\\'s the difference between 2.4GHz and 5GHz?')">
                            📊 2.4GHz vs 5GHz
                        </button>
                    </div>
                </div>

                <div class="category">
                    <h4>🔧 Setup & Configuration</h4>
                    <div class="question-buttons">
                        <button onclick="useSuggestedQuestion('How do I change my WiFi password?')">
                            🔑 Change WiFi password
                        </button>
                        <button onclick="useSuggestedQuestion('How do I reset my router to factory settings?')">
                            🔄 Factory reset
                        </button>
                        <button onclick="useSuggestedQuestion('How do I setup parental controls?')">
                            👪 Parental controls
                        </button>
                    </div>
                </div>

                <div class="category">
                    <h4>📱 Router Models</h4>
                    <div class="question-buttons">
                        <button onclick="useSuggestedQuestion('Tell me about Huawei A5 router')">
                            📱 Huawei A5 info
                        </button>
                        <button onclick="useSuggestedQuestion('Tell me about Huawei V5 router')">
                            🚀 Huawei V5 info
                        </button>
                        <button onclick="useSuggestedQuestion('What are the differences between A5 and V5?')">
                            ⚖️ Compare A5 vs V5
                        </button>
                    </div>
                </div>
            </div>

            <!-- View More Questions Toggle -->
            <div class="view-more">
                <button onclick="toggleAllQuestions()" id="toggleQuestionsBtn" class="toggle-btn">
                    ▼ Show All Questions ({{ suggested_questions|length }} total)
                </button>
            </div>

            <div id="allQuestions" class="all-questions" style="display: none;">
                <h4>📋 Complete Question List:</h4>
                <div class="question-grid">
                    {% for question in suggested_questions %}
                        <button onclick="useSuggestedQuestion('{{ question|escapejs }}')" class="question-item">
                            {{ question }}
                        </button>
                    {% endfor %}
                </div>
            </div>
        </div>

        <!-- Footer -->
        <div class="footer">
            <p>⚠️ This is an automated support assistant. For urgent issues, please <a href="/contact">contact support directly</a>.</p>
            <p class="small">Supporting Huawei A5 (2.4GHz) and Huawei V5 (Dual Band) routers</p>
        </div>
    </div>

    <!-- JavaScript Section -->
    <script>
        // Scroll to bottom of chat on load
        window.onload = function() {
            scrollToBottom();
            document.getElementById('userInput').focus();
        };

        // Auto-resize input as user types
        const input = document.getElementById('userInput');
        input.addEventListener('input', function() {
            this.style.height = 'auto';
            this.style.height = (this.scrollHeight) + 'px';
        });

        function scrollToBottom() {
            const chatMessages = document.getElementById('chatMessages');
            chatMessages.scrollTop = chatMessages.scrollHeight;
        }

        function handleKeyPress(event) {
            if (event.key === 'Enter' && !event.shiftKey) {
                event.preventDefault(); // Prevent new line
                sendMessage();
            }
        }

        function useSuggestedQuestion(question) {
            document.getElementById('userInput').value = question;
            // Adjust input height
            input.style.height = 'auto';
            input.style.height = (input.scrollHeight) + 'px';
            sendMessage();
        }

        function showTypingIndicator() {
            document.getElementById('typingIndicator').style.display = 'block';
            scrollToBottom();
        }

        function hideTypingIndicator() {
            document.getElementById('typingIndicator').style.display = 'none';
        }

        function sendMessage() {
            const input = document.getElementById('userInput');
            const message = input.value.trim();

            if (message === '') {
                // Shake animation for empty input
                input.style.borderColor = '#ff4757';
                setTimeout(() => {
                    input.style.borderColor = '#ddd';
                }, 1000);
                return;
            }

            // Disable input while sending
            input.disabled = true;
            document.getElementById('sendButton').disabled = true;

            // Display user message immediately
            addMessageToChat('user', message);

            // Clear input and reset height
            input.value = '';
            input.style.height = 'auto';

            // Show typing indicator
            showTypingIndicator();

            // Send to server
            fetch('/get-response/', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json',
                    'X-CSRFToken': getCookie('csrftoken')
                },
                body: JSON.stringify({message: message})
            })
            .then(response => {
                if (!response.ok) {
                    throw new Error('Network response was not ok');
                }
                return response.json();
            })
            .then(data => {
                // Hide typing indicator
                hideTypingIndicator();

                // Display bot response with slight delay for natural feel
                setTimeout(() => {
                    addMessageToChat('bot', data.response);

                    // Re-enable input
                    input.disabled = false;
                    document.getElementById('sendButton').disabled = false;
                    input.focus();
                }, 500);
            })
            .catch(error => {
                console.error('Error:', error);
                hideTypingIndicator();
                addMessageToChat('bot', '😔 Sorry, I encountered an error. Please try again or contact support if the issue persists.');

                // Re-enable input
                input.disabled = false;
                document.getElementById('sendButton').disabled = false;
                input.focus();
            });
        }

        function addMessageToChat(type, text) {
            const chatMessages = document.getElementById('chatMessages');
            const messageDiv = document.createElement('div');
            messageDiv.className = `message ${type}-message`;

            // Add timestamp
            const now = new Date();
            const timeStr = now.toLocaleTimeString([], { hour: '2-digit', minute: '2-digit' });

            messageDiv.innerHTML = `
                <div class="message-content">
                    <span class="icon">${type === 'bot' ? '🤖' : '👤'}</span>
                    ${text.replace(/\\n/g, '<br>')}
                </div>
                <div class="message-time">${timeStr}</div>
            `;

            chatMessages.appendChild(messageDiv);
            scrollToBottom();

            // Highlight new message briefly
            messageDiv.style.animation = 'fadeIn 0.3s';
        }

        function clearHistory() {
            // Create custom confirmation modal
            const confirmModal = document.createElement('div');
            confirmModal.className = 'confirm-modal';
            confirmModal.innerHTML = `
                <div class="confirm-modal-content">
                    <div class="confirm-modal-header">
                        <span class="confirm-icon">🗑️</span>
                        <h3>Clear Chat History?</h3>
                    </div>
                    <div class="confirm-modal-body">
                        <p>Are you sure you want to clear all conversation history?</p>
                        <p class="warning-text">This action cannot be undone.</p>
                    </div>
                    <div class="confirm-modal-footer">
                        <button class="confirm-btn cancel-btn" onclick="closeConfirmModal()">No, Keep History</button>
                        <button class="confirm-btn confirm-delete-btn" onclick="proceedClearHistory()">Yes, Clear History</button>
                    </div>
                </div>
            `;

            document.body.appendChild(confirmModal);

            // Add escape key handler
            document.addEventListener('keydown', escapeKeyHandler);
        }

        function closeConfirmModal() {
            const modal = document.querySelector('.confirm-modal');
            if (modal) {
                modal.remove();
            }
            document.removeEventListener('keydown', escapeKeyHandler);
        }

        function escapeKeyHandler(e) {
            if (e.key === 'Escape') {
                closeConfirmModal();
            }
        }

        function proceedClearHistory() {
            // Close the modal first for better UX
            closeConfirmModal();

            // Show loading state on the clear button
            const clearBtn = document.querySelector('.clear-btn');
            const originalText = clearBtn.innerHTML;
            clearBtn.innerHTML = '⏰ Clearing...';
            clearBtn.disabled = true;

            // Send request to clear history
            fetch('/clear-history/', {
                method: 'POST',
                headers: {
                    'X-CSRFToken': getCookie('csrftoken')
                }
            })
            .then(response => {
                if (!response.ok) {
                    throw new Error('Failed to clear history');
                }
                return response.json();
            })
            .then(data => {
                if (data.success) {
                    // Show success message
                    showNotification('Chat history cleared successfully!', 'success');

                    // Reload after a brief delay to show success message
                    setTimeout(() => {
                        location.reload();
                    }, 1000);
                } else {
                    throw new Error('Server returned unsuccessful response');
                }
            })
            .catch(error => {
                console.error('Error clearing history:', error);

                // Show error message
                showNotification('Failed to clear history. Please try again.', 'error');

                // Reset button
                clearBtn.innerHTML = originalText;
                clearBtn.disabled = false;
            });
        }

        function showNotification(message, type) {
            // Create notification element
            const notification = document.createElement('div');
            notification.className = `notification ${type}`;
            notification.innerHTML = `
                <div class="notification-content">
                    <span class="notification-icon">${type === 'success' ? '✅' : '❌'}</span>
                    <span class="notification-message">${message}</span>
                </div>
            `;

            document.body.appendChild(notification);

            // Auto-remove after 3 seconds
            setTimeout(() => {
                notification.style.animation = 'slideOut 0.3s ease';
                setTimeout(() => {
                    notification.remove();
                }, 300);
            }, 3000);
        }

        function toggleAllQuestions() {
            const allQuestions = document.getElementById('allQuestions');
            const toggleBtn = document.getElementById('toggleQuestionsBtn');

            if (allQuestions.style.display === 'none') {
                allQuestions.style.display = 'block';
                toggleBtn.innerHTML = '▲ Hide Questions';
            } else {
                allQuestions.style.display = 'none';
                toggleBtn.innerHTML = '▼ Show All Questions ({{ suggested_questions|length }} total)';
            }
        }

        // Function to get CSRF token
        function getCookie(name) {
            let cookieValue = null;
            if (document.cookie && document.cookie !== '') {
                const cookies = document.cookie.split(';');
                for (let i = 0; i < cookies.length; i++) {
                    const cookie = cookies[i].trim();
                    if (cookie.substring(0, name.length + 1) === (name + '=')) {
                        cookieValue = decodeURIComponent(cookie.substring(name.length + 1));
                        break;
                    }
                }
            }
            return cookieValue;
        }

        // Add keyboard shortcut (Ctrl+Enter to send)
        document.addEventListener('keydown', function(e) {
            if (e.ctrlKey && e.key === 'Enter') {
                sendMessage();
            }
        });

        // Check connection status periodically
        function checkConnection() {
            if (!navigator.onLine) {
                addMessageToChat('bot', '⚠️ You appear to be offline. Please check your internet connection.');
            }
        }

        window.addEventListener('offline', checkConnection);

        // Add animation keyframes dynamically
        const style = document.createElement('style');
        style.textContent = `
            @keyframes fadeIn {
                from { opacity: 0; transform: translateY(10px); }
                to { opacity: 1; transform: translateY(0); }
            }

            .typing-indicator .dot {
                display: inline-block;
                width: 8px;
                height: 8px;
                border-radius: 50%;
                background-color: #666;
                margin: 0 2px;
                animation: typing 1.4s infinite;
            }

            .typing-indicator .dot:nth-child(2) { animation-delay: 0.2s; }
            .typing-indicator .dot:nth-child(3) { animation-delay: 0.4s; }

            @keyframes typing {
                0%, 60%, 100% { transform: translateY(0); opacity: 0.4; }
                30% { transform: translateY(-10px); opacity: 1; }
            }

            .message-time {
                font-size: 10px;
                color: #999;
                margin-top: 4px;
                text-align: right;
            }

            .quick-actions {
                display: flex;
                gap: 8px;
                padding: 10px 15px;
                background: #f5f5f5;
                border-top: 1px solid #eee;
                overflow-x: auto;
                white-space: nowrap;
            }

            .quick-btn {
                background: white;
                border: 1px solid #ddd;
                border-radius: 20px;
                padding: 6px 12px;
                font-size: 12px;
                cursor: pointer;
                transition: all 0.3s;
                margin: 0;
                color: #666;
            }

            .quick-btn:hover {
                background: #667eea;
                color: white;
                border-color: #667eea;
            }

            .category {
                margin-bottom: 20px;
            }

            .category h4 {
                color: #666;
                font-size: 14px;
                margin-bottom: 10px;
                padding-left: 5px;
                border-left: 3px solid #667eea;
            }

            .all-questions {
                margin-top: 20px;
                padding-top: 20px;
                border-top: 1px solid #ddd;
            }

            .question-grid {
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(250px, 1fr));
                gap: 10px;
                margin-top: 15px;
            }

            .question-item {
                background: white;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 10px;
                text-align: left;
                font-size: 13px;
                color: #333;
                transition: all 0.3s;
                margin: 0;
            }

            .question-item:hover {
                background: #f0f4ff;
                border-color: #667eea;
                transform: translateY(-2px);
                box-shadow: 0 4px 8px rgba(102, 126, 234, 0.2);
            }

            .toggle-btn {
                background: none;
                border: 1px solid #ddd;
                color: #667eea;
                padding: 8px 16px;
                font-size: 13px;
                cursor: pointer;
                margin: 10px 0;
                border-radius: 5px;
            }

            .toggle-btn:hover {
                background: #667eea;
                color: white;
                border-color: #667eea;
            }

            .footer {
                background: #f5f5f5;
                padding: 15px;
                text-align: center;
                border-top: 1px solid #ddd;
                font-size: 12px;
                color: #666;
            }

            .footer .small {
                font-size: 11px;
                color: #999;
                margin-top: 5px;
            }

            .header {
                position: relative;
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 20px;
            }

            .clear-btn {
                position: absolute;
                top: 20px;
                right: 20px;
                background: rgba(255, 255, 255, 0.2);
                border: 1px solid rgba(255, 255, 255, 0.5);
                color: white;
                padding: 5px 10px;
                border-radius: 5px;
                font-size: 12px;
                cursor: pointer;
                transition: all 0.3s;
            }

            .clear-btn:hover {
                background: rgba(255, 255, 255, 0.3);
                transform: scale(1.05);
            }

            .shortcut {
                background: rgba(255, 255, 255, 0.2);
                padding: 2px 6px;
                border-radius: 4px;
                font-size: 10px;
                margin-left: 5px;
            }

            /* Mobile responsiveness */
            @media (max-width: 600px) {
                .container {
                    margin: 0;
                    border-radius: 0;
                }

                .question-grid {
                    grid-template-columns: 1fr;
                }

                .quick-actions {
                    flex-wrap: wrap;
                }

                .quick-btn {
                    flex: 1 1 auto;
                }
            }
        `;
        document.head.appendChild(style);
    </script>
</body>
</html>'''

        self.add_code_block(chat_html, "html")

        self.doc.add_heading('5.2 HTML Template Explanation', 3)

        html_sections = [
            ('Meta Tags', 'Viewport settings for mobile responsiveness, theme color for PWA support'),
            ('Header Section', 'Contains title and clear chat button with gradient background'),
            ('Chat Messages Container', 'Displays conversation history with bot/user message styling'),
            ('Input Area', 'Text input with send button, auto-resize functionality'),
            ('Quick Actions', 'Three common questions for quick access'),
            ('Suggested Questions', 'Categorized question buttons for easy access'),
            ('JavaScript Functions', 'Handles message sending, UI updates, and API calls')
        ]

        for section, desc in html_sections:
            p = self.doc.add_paragraph()
            p.add_run(f'{section}: ').bold = True
            p.add_run(desc)

        # contact.html
        self.doc.add_heading('FILE: chatbot/templates/contact.html', 2)
        contact_html = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Contact Support</title>
    {% load static %}
    <link rel="stylesheet" href="{% static 'chatbot/style.css' %}">
</head>
<body>
    <div class="container">
        <h1>📞 Contact Support</h1>
        <div style="padding: 40px; text-align: center;">
            <p style="margin-bottom: 20px;">Our support team is here to help you!</p>

            <div style="margin: 30px 0;">
                <h3 style="color: #667eea; margin-bottom: 10px;">📱 Phone Support</h3>
                <p style="font-size: 20px;">1-800-ROUTER-HELP</p>
                <p style="color: #666;">Mon-Fri: 8am - 8pm</p>
            </div>

            <div style="margin: 30px 0;">
                <h3 style="color: #667eea; margin-bottom: 10px;">✉️ Email Support</h3>
                <p>support@routerhelp.com</p>
                <p style="color: #666;">We'll respond within 24 hours</p>
            </div>

            <a href="/" style="display: inline-block; margin-top: 20px; color: #667eea; text-decoration: none;">
                ← Back to Chat
            </a>
        </div>
    </div>
</body>
</html>'''

        self.add_code_block(contact_html, "html")

        self.doc.add_page_break()

    def chapter6_static_files(self):
        """Chapter 6: Static Files and Styling"""
        self.add_heading_centered('CHAPTER 6', 1)
        self.add_heading_centered('STATIC FILES AND STYLING', 2)

        # style.css
        self.doc.add_heading('FILE: chatbot/static/chatbot/style.css', 2)

        css_code = '''/* Mobile-First CSS for Router Support Chatbot */
:root {
    --primary-gradient: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    --primary-color: #667eea;
    --secondary-color: #764ba2;
    --danger-color: #ff4757;
    --success-color: #00c853;
    --warning-color: #ffa502;
    --text-dark: #333;
    --text-light: #666;
    --text-lighter: #999;
    --bg-light: #f9f9f9;
    --bg-white: #ffffff;
    --border-color: #ddd;
    --shadow-sm: 0 2px 8px rgba(0,0,0,0.1);
    --shadow-md: 0 4px 12px rgba(0,0,0,0.15);
    --shadow-lg: 0 10px 30px rgba(0,0,0,0.2);
    --radius-sm: 8px;
    --radius-md: 12px;
    --radius-lg: 20px;
    --spacing-xs: 4px;
    --spacing-sm: 8px;
    --spacing-md: 15px;
    --spacing-lg: 20px;
    --spacing-xl: 30px;
}

* {
    margin: 0;
    padding: 0;
    box-sizing: border-box;
    -webkit-tap-highlight-color: transparent;
}

body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
    background: var(--primary-gradient);
    min-height: 100vh;
    display: flex;
    align-items: center;
    justify-content: center;
    padding: 0;
    margin: 0;
}

/* Mobile-First Container */
.container {
    width: 100%;
    max-width: 100%;
    min-height: 100vh;
    background: var(--bg-white);
    display: flex;
    flex-direction: column;
    box-shadow: none;
    border-radius: 0;
}

/* Header - Mobile Optimized */
.header {
    position: relative;
    background: var(--primary-gradient);
    color: white;
    padding: var(--spacing-lg);
    padding-top: calc(env(safe-area-inset-top) + var(--spacing-lg));
    text-align: center;
}

h1 {
    font-size: 1.5rem;
    margin: 0;
    font-weight: 600;
    letter-spacing: -0.5px;
    padding-right: 50px;
}

.subtitle {
    font-size: 0.85rem;
    opacity: 0.9;
    margin-top: var(--spacing-xs);
    padding: 0 var(--spacing-lg);
    background: rgba(255,255,255,0.1);
    padding: var(--spacing-sm) var(--spacing-md);
    border-radius: var(--radius-lg);
    display: inline-block;
    margin-top: var(--spacing-sm);
}

/* Clear Button - Mobile Optimized */
.clear-btn {
    position: absolute;
    top: calc(env(safe-area-inset-top) + var(--spacing-md));
    right: var(--spacing-md);
    background: rgba(255, 255, 255, 0.2);
    border: 1px solid rgba(255, 255, 255, 0.3);
    color: white;
    padding: var(--spacing-sm) var(--spacing-sm);
    border-radius: var(--radius-lg);
    font-size: 0.75rem;
    cursor: pointer;
    backdrop-filter: blur(5px);
    -webkit-backdrop-filter: blur(5px);
    display: flex;
    align-items: center;
    gap: 4px;
    min-width: 70px;
    justify-content: center;
}

.clear-btn:active {
    background: rgba(255, 255, 255, 0.3);
    transform: scale(0.95);
}

/* Chat Container - Mobile Optimized */
.chat-container {
    flex: 1;
    display: flex;
    flex-direction: column;
    background: var(--bg-light);
    height: calc(100vh - 180px);
    min-height: 400px;
}

/* Chat Messages Area */
.chat-messages {
    flex: 1;
    overflow-y: auto;
    padding: var(--spacing-md);
    -webkit-overflow-scrolling: touch;
    scroll-behavior: smooth;
    display: flex;
    flex-direction: column;
    gap: var(--spacing-sm);
}

/* Message Bubbles - Mobile Optimized */
.message {
    max-width: 85%;
    animation: fadeIn 0.3s ease;
}

.bot-message {
    align-self: flex-start;
}

.user-message {
    align-self: flex-end;
}

.message-content {
    padding: var(--spacing-md);
    border-radius: var(--radius-lg);
    font-size: 0.9rem;
    line-height: 1.4;
    word-wrap: break-word;
    box-shadow: var(--shadow-sm);
    position: relative;
}

.bot-message .message-content {
    background: var(--bg-white);
    border: 1px solid var(--border-color);
    border-top-left-radius: 5px;
    color: var(--text-dark);
}

.user-message .message-content {
    background: var(--primary-gradient);
    color: white;
    border-top-right-radius: 5px;
}

.message-content a {
    color: inherit;
    text-decoration: underline;
    text-decoration-color: rgba(255,255,255,0.5);
}

.icon {
    margin-right: var(--spacing-sm);
    font-size: 1.1rem;
    display: inline-block;
}

/* Message Time */
.message-time {
    font-size: 0.65rem;
    color: var(--text-lighter);
    margin-top: 2px;
    padding: 0 var(--spacing-xs);
}

.user-message .message-time {
    text-align: right;
}

/* Input Area - Mobile Optimized */
.input-area {
    background: var(--bg-white);
    padding: var(--spacing-sm);
    border-top: 1px solid var(--border-color);
    display: flex;
    gap: var(--spacing-sm);
    align-items: flex-end;
    position: sticky;
    bottom: 0;
    padding-bottom: max(env(safe-area-inset-bottom), var(--spacing-sm));
}

#userInput {
    flex: 1;
    padding: var(--spacing-md);
    border: 1px solid var(--border-color);
    border-radius: var(--radius-lg);
    font-size: 0.95rem;
    outline: none;
    transition: all 0.3s;
    max-height: 120px;
    resize: none;
    font-family: inherit;
    background: var(--bg-white);
    -webkit-appearance: none;
}

#userInput:focus {
    border-color: var(--primary-color);
    box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
}

#userInput:disabled {
    background: #f5f5f5;
    color: var(--text-lighter);
}

#sendButton {
    background: var(--primary-gradient);
    color: white;
    border: none;
    border-radius: var(--radius-lg);
    padding: var(--spacing-md) var(--spacing-lg);
    font-size: 0.95rem;
    font-weight: 500;
    cursor: pointer;
    transition: all 0.3s;
    display: flex;
    align-items: center;
    gap: var(--spacing-xs);
    min-width: 70px;
    justify-content: center;
    -webkit-appearance: none;
}

#sendButton:active {
    transform: scale(0.95);
}

#sendButton:disabled {
    opacity: 0.5;
    pointer-events: none;
}

/* Quick Actions - Mobile Optimized */
.quick-actions {
    display: flex;
    gap: var(--spacing-sm);
    padding: var(--spacing-sm);
    background: var(--bg-white);
    overflow-x: auto;
    -webkit-overflow-scrolling: touch;
    scrollbar-width: none;
    white-space: nowrap;
    border-top: 1px solid var(--border-color);
}

.quick-actions::-webkit-scrollbar {
    display: none;
}

.quick-btn {
    background: var(--bg-light);
    border: 1px solid var(--border-color);
    border-radius: var(--radius-lg);
    padding: var(--spacing-sm) var(--spacing-md);
    font-size: 0.8rem;
    cursor: pointer;
    transition: all 0.3s;
    color: var(--text-dark);
    flex: 0 0 auto;
    -webkit-appearance: none;
}

.quick-btn:active {
    background: var(--primary-color);
    color: white;
    border-color: var(--primary-color);
    transform: scale(0.95);
}

/* Suggested Questions - Mobile Optimized */
.suggested-questions {
    background: var(--bg-white);
    padding: var(--spacing-md);
    border-top: 1px solid var(--border-color);
    max-height: 40vh;
    overflow-y: auto;
    -webkit-overflow-scrolling: touch;
}

.suggested-questions h3 {
    font-size: 1rem;
    color: var(--text-dark);
    margin-bottom: var(--spacing-md);
    display: flex;
    align-items: center;
    gap: var(--spacing-xs);
}

.category {
    margin-bottom: var(--spacing-lg);
}

.category h4 {
    color: var(--primary-color);
    font-size: 0.9rem;
    margin-bottom: var(--spacing-sm);
    padding-left: var(--spacing-sm);
    border-left: 3px solid var(--primary-color);
}

.question-buttons {
    display: flex;
    flex-wrap: wrap;
    gap: var(--spacing-sm);
}

.question-buttons button {
    background: var(--bg-light);
    border: 1px solid var(--border-color);
    border-radius: var(--radius-lg);
    padding: var(--spacing-sm) var(--spacing-md);
    font-size: 0.8rem;
    cursor: pointer;
    transition: all 0.3s;
    color: var(--text-dark);
    flex: 1 1 calc(50% - var(--spacing-sm));
    min-width: 120px;
    text-align: left;
    display: flex;
    align-items: center;
    gap: var(--spacing-xs);
    -webkit-appearance: none;
}

.question-buttons button:active {
    background: var(--primary-gradient);
    color: white;
    border-color: var(--primary-color);
    transform: translateY(-2px);
}

/* All Questions Grid */
.all-questions {
    margin-top: var(--spacing-lg);
    padding-top: var(--spacing-lg);
    border-top: 2px solid var(--border-color);
}

.question-grid {
    display: grid;
    grid-template-columns: repeat(auto-fill, minmax(140px, 1fr));
    gap: var(--spacing-sm);
    margin-top: var(--spacing-sm);
}

.question-item {
    background: var(--bg-white);
    border: 1px solid var(--border-color);
    border-radius: var(--radius-md);
    padding: var(--spacing-sm);
    font-size: 0.75rem;
    text-align: left;
    color: var(--text-dark);
    cursor: pointer;
    transition: all 0.3s;
    line-height: 1.3;
    min-height: 60px;
    word-break: break-word;
    -webkit-appearance: none;
}

.question-item:active {
    background: var(--primary-gradient);
    color: white;
    border-color: var(--primary-color);
    transform: translateY(-2px);
}

/* Toggle Button */
.toggle-btn {
    background: none;
    border: 1px solid var(--border-color);
    color: var(--primary-color);
    padding: var(--spacing-sm) var(--spacing-md);
    font-size: 0.8rem;
    cursor: pointer;
    border-radius: var(--radius-lg);
    width: 100%;
    margin: var(--spacing-sm) 0;
    -webkit-appearance: none;
}

.toggle-btn:active {
    background: var(--primary-gradient);
    color: white;
    border-color: var(--primary-color);
}

/* Footer */
.footer {
    background: var(--bg-light);
    padding: var(--spacing-md);
    text-align: center;
    border-top: 1px solid var(--border-color);
    font-size: 0.75rem;
    color: var(--text-light);
    padding-bottom: max(env(safe-area-inset-bottom), var(--spacing-md));
}

.footer a {
    color: var(--primary-color);
    text-decoration: none;
    font-weight: 500;
}

.footer .small {
    font-size: 0.65rem;
    color: var(--text-lighter);
    margin-top: var(--spacing-xs);
}

/* Typing Indicator */
.typing-indicator .message-content {
    display: flex;
    align-items: center;
    gap: var(--spacing-xs);
}

.typing-indicator .dot {
    width: 8px;
    height: 8px;
    border-radius: 50%;
    background-color: var(--text-lighter);
    animation: typing 1.4s infinite;
    display: inline-block;
}

.typing-indicator .dot:nth-child(2) { animation-delay: 0.2s; }
.typing-indicator .dot:nth-child(3) { animation-delay: 0.4s; }

/* Confirmation Modal - Mobile Optimized */
.confirm-modal {
    position: fixed;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    background: rgba(0, 0, 0, 0.5);
    display: flex;
    justify-content: center;
    align-items: center;
    z-index: 1000;
    animation: fadeIn 0.3s ease;
    padding: var(--spacing-md);
    backdrop-filter: blur(5px);
    -webkit-backdrop-filter: blur(5px);
}

.confirm-modal-content {
    background: white;
    border-radius: var(--radius-lg);
    width: 100%;
    max-width: 320px;
    box-shadow: var(--shadow-lg);
    animation: slideUp 0.3s ease;
    overflow: hidden;
}

.confirm-modal-header {
    background: var(--primary-gradient);
    color: white;
    padding: var(--spacing-lg);
    display: flex;
    align-items: center;
    gap: var(--spacing-sm);
}

.confirm-icon {
    font-size: 1.5rem;
}

.confirm-modal-header h3 {
    margin: 0;
    font-size: 1.1rem;
    font-weight: 500;
}

.confirm-modal-body {
    padding: var(--spacing-lg);
    text-align: center;
}

.confirm-modal-body p {
    margin: var(--spacing-sm) 0;
    color: var(--text-dark);
    font-size: 0.95rem;
}

.warning-text {
    color: var(--danger-color) !important;
    font-size: 0.8rem !important;
    font-weight: 500;
}

.confirm-modal-footer {
    padding: var(--spacing-md);
    display: flex;
    gap: var(--spacing-sm);
    background: var(--bg-light);
    border-top: 1px solid var(--border-color);
}

.confirm-btn {
    flex: 1;
    padding: var(--spacing-md);
    border: none;
    border-radius: var(--radius-md);
    font-size: 0.9rem;
    cursor: pointer;
    transition: all 0.3s;
    font-weight: 500;
    -webkit-appearance: none;
}

.cancel-btn {
    background: white;
    color: var(--text-dark);
    border: 1px solid var(--border-color);
}

.cancel-btn:active {
    background: #f0f0f0;
    transform: scale(0.98);
}

.confirm-delete-btn {
    background: var(--danger-color);
    color: white;
}

.confirm-delete-btn:active {
    background: #ff3344;
    transform: scale(0.98);
}

/* Notifications - Mobile Optimized */
.notification {
    position: fixed;
    top: max(env(safe-area-inset-top), 10px);
    left: var(--spacing-md);
    right: var(--spacing-md);
    z-index: 1001;
    animation: slideDown 0.3s ease;
    max-width: 400px;
    margin: 0 auto;
}

.notification-content {
    background: white;
    border-radius: var(--radius-md);
    padding: var(--spacing-md);
    box-shadow: var(--shadow-lg);
    display: flex;
    align-items: center;
    gap: var(--spacing-sm);
    border-left: 4px solid;
}

.notification.success .notification-content {
    border-left-color: var(--success-color);
}

.notification.error .notification-content {
    border-left-color: var(--danger-color);
}

.notification-icon {
    font-size: 1.2rem;
}

.notification-message {
    color: var(--text-dark);
    font-size: 0.9rem;
    flex: 1;
}

/* Animations */
@keyframes fadeIn {
    from { opacity: 0; }
    to { opacity: 1; }
}

@keyframes slideUp {
    from {
        transform: translateY(50px);
        opacity: 0;
    }
    to {
        transform: translateY(0);
        opacity: 1;
    }
}

@keyframes slideDown {
    from {
        transform: translateY(-100%);
        opacity: 0;
    }
    to {
        transform: translateY(0);
        opacity: 1;
    }
}

@keyframes slideOut {
    from {
        transform: translateX(0);
        opacity: 1;
    }
    to {
        transform: translateX(100%);
        opacity: 0;
    }
}

@keyframes typing {
    0%, 60%, 100% { transform: translateY(0); opacity: 0.4; }
    30% { transform: translateY(-10px); opacity: 1; }
}

/* Tablet and Desktop Improvements */
@media (min-width: 768px) {
    body {
        padding: var(--spacing-lg);
    }

    .container {
        max-width: 800px;
        min-height: auto;
        border-radius: var(--radius-lg);
        box-shadow: var(--shadow-lg);
    }

    h1 {
        font-size: 1.8rem;
        padding-right: 0;
    }

    .chat-container {
        height: 500px;
    }

    .quick-actions {
        justify-content: center;
    }

    .question-buttons button {
        flex: 0 1 auto;
        min-width: 180px;
    }

    .question-grid {
        grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
    }

    .confirm-modal-content {
        max-width: 400px;
    }

    .notification {
        left: auto;
        right: var(--spacing-lg);
        min-width: 300px;
    }
}

/* Landscape Mode Optimization */
@media (max-height: 600px) and (orientation: landscape) {
    .chat-container {
        height: 60vh;
    }

    .suggested-questions {
        max-height: 30vh;
    }

    .confirm-modal-content {
        max-height: 90vh;
        overflow-y: auto;
    }
}

/* Dark Mode Support (Optional) */
@media (prefers-color-scheme: dark) {
    :root {
        --bg-white: #1a1a1a;
        --bg-light: #2d2d2d;
        --text-dark: #ffffff;
        --text-light: #e0e0e0;
        --text-lighter: #a0a0a0;
        --border-color: #404040;
    }

    body {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
    }

    .message-content {
        box-shadow: 0 2px 8px rgba(0,0,0,0.3);
    }

    #userInput {
        background: #2d2d2d;
        color: white;
        border-color: #404040;
    }

    #userInput::placeholder {
        color: #808080;
    }
}

/* Safe Area Support for Notched Phones */
@supports (padding: max(0px)) {
    .header {
        padding-top: max(var(--spacing-lg), env(safe-area-inset-top));
    }

    .input-area {
        padding-bottom: max(var(--spacing-sm), env(safe-area-inset-bottom));
    }

    .footer {
        padding-bottom: max(var(--spacing-md), env(safe-area-inset-bottom));
    }
}

/* Touch-friendly improvements */
button, 
.question-buttons button,
.quick-btn,
.question-item,
.toggle-btn {
    touch-action: manipulation;
    cursor: pointer;
    -webkit-tap-highlight-color: transparent;
}

/* Loading state */
.loading {
    opacity: 0.7;
    pointer-events: none;
}

/* Error state */
.error-shake {
    animation: shake 0.5s ease;
}

@keyframes shake {
    0%, 100% { transform: translateX(0); }
    25% { transform: translateX(-5px); }
    75% { transform: translateX(5px); }
}

/* Print styles (for support tickets) */
@media print {
    .input-area,
    .quick-actions,
    .suggested-questions,
    .clear-btn,
    #sendButton {
        display: none !important;
    }

    .chat-container {
        height: auto;
        overflow: visible;
    }

    .message {
        break-inside: avoid;
    }
}'''

        self.add_code_block(css_code, "css")

        self.doc.add_heading('6.1 CSS Architecture Explanation', 3)

        css_features = [
            ('CSS Variables (Custom Properties)', 'Defines reusable design tokens for consistent theming'),
            ('Mobile-First Approach', 'Base styles for mobile with progressive enhancement for larger screens'),
            ('CSS Grid and Flexbox', 'Modern layout techniques for responsive design'),
            ('CSS Animations', 'Smooth transitions for typing indicator and modals'),
            ('Media Queries', 'Responsive breakpoints for tablet and desktop optimization'),
            ('Dark Mode Support', 'Automatic theme switching based on user preference'),
            ('Safe Area Support', 'Proper padding for notched phones using env() variables'),
            ('Touch Optimization', 'Larger tap targets and removed tap highlights')
        ]

        for feature, desc in css_features:
            p = self.doc.add_paragraph()
            p.add_run(f'{feature}: ').bold = True
            p.add_run(desc)

        # manifest.json
        self.doc.add_heading('FILE: chatbot/static/chatbot/manifest.json', 2)
        manifest_code = '''{
  "short_name": "Router Support",
  "name": "Router Support Chatbot",
  "icons": [
    {
      "src": "icon-192.png",
      "type": "image/png",
      "sizes": "192x192"
    },
    {
      "src": "icon-512.png",
      "type": "image/png",
      "sizes": "512x512"
    }
  ],
  "start_url": "/",
  "display": "standalone",
  "theme_color": "#667eea",
  "background_color": "#ffffff",
  "orientation": "portrait"
}'''

        self.add_code_block(manifest_code, "json")

        self.doc.add_paragraph(
            'The manifest.json file enables Progressive Web App (PWA) features, allowing users to install the '
            'chatbot on their home screen and use it like a native app.'
        )

        self.doc.add_page_break()

    def chapter7_deployment(self):
        """Chapter 7: Deployment Configuration"""
        self.add_heading_centered('CHAPTER 7', 1)
        self.add_heading_centered('DEPLOYMENT CONFIGURATION', 2)

        # requirements.txt
        self.doc.add_heading('FILE: requirements.txt', 2)
        requirements = '''Django>=4.2.0,<5.0.0
gunicorn==20.1.0
whitenoise==6.4.0
psycopg2-binary==2.9.9
python-dotenv==1.0.0
psutil==5.9.5'''

        self.add_code_block(requirements, "txt")

        self.doc.add_heading('7.1 Requirements Explanation', 3)
        req_explanations = [
            ('Django', 'Web framework for building the application'),
            ('gunicorn', 'WSGI HTTP server for running Django in production'),
            ('whitenoise', 'Serves static files efficiently in production'),
            ('psycopg2-binary', 'PostgreSQL database adapter'),
            ('python-dotenv', 'Loads environment variables from .env file'),
            ('psutil', 'System monitoring for health checks')
        ]

        for pkg, desc in req_explanations:
            p = self.doc.add_paragraph()
            p.add_run(f'{pkg}: ').bold = True
            p.add_run(desc)

        # runtime.txt
        self.doc.add_heading('FILE: runtime.txt', 2)
        self.add_code_block('python-3.11.0', "txt")

        # Procfile
        self.doc.add_heading('FILE: Procfile', 2)
        self.add_code_block('web: gunicorn router_chatbot.wsgi --log-file -', "bash")

        # render.yaml
        self.doc.add_heading('FILE: render.yaml', 2)
        render_yaml = '''services:
  - type: web
    name: router-support-chatbot
    runtime: python
    buildCommand: pip install -r requirements.txt && python manage.py collectstatic --noinput
    startCommand: gunicorn router_chatbot.wsgi --log-file -
    envVars:
      - key: SECRET_KEY
        generateValue: true
      - key: DEBUG
        value: false
      - key: PYTHON_VERSION
        value: 3.11.0'''

        self.add_code_block(render_yaml, "yaml")

        # build.sh
        self.doc.add_heading('FILE: build.sh', 2)
        build_sh = '''#!/usr/bin/env bash
# Exit on error
set -o errexit

# Install dependencies
pip install -r requirements.txt

# Run migrations
python manage.py makemigrations
python manage.py migrate

# Collect static files
python manage.py collectstatic --no-input'''

        self.add_code_block(build_sh, "bash")

        self.doc.add_heading('7.2 Deployment Process', 3)

        deploy_steps = [
            ('1. Environment Setup', 'Create virtual environment and install dependencies'),
            ('2. Git Repository', 'Initialize git and push to GitHub'),
            ('3. Render Deployment', 'Connect GitHub repository to Render and configure service'),
            ('4. Environment Variables', 'Set SECRET_KEY, DEBUG=False, and DATABASE_URL'),
            ('5. Database Setup', 'Provision PostgreSQL database and connect to web service'),
            ('6. Static Files', 'Whitenoise serves compressed and cached static files'),
            ('7. Monitoring Setup', 'Configure UptimeRobot to ping /ping/ endpoint every 5 minutes')
        ]

        for step, desc in deploy_steps:
            p = self.doc.add_paragraph()
            p.add_run(f'{step}: ').bold = True
            p.add_run(desc)

        self.doc.add_page_break()

    def chapter8_features(self):
        """Chapter 8: Features Summary"""
        self.add_heading_centered('CHAPTER 8', 1)
        self.add_heading_centered('FEATURES SUMMARY', 2)

        self.doc.add_heading('8.1 System Features', 3)

        features = [
            ('Rule-Based Chatbot', 'Intelligent keyword matching with 25+ intents and 100+ keywords'),
            ('Session Management', 'Persists conversation history across page reloads'),
            ('Mobile-First Design', 'Fully responsive interface optimized for all devices'),
            ('Progressive Web App', 'Installable on home screen with native app experience'),
            ('Quick Action Buttons', 'One-click access to common questions'),
            ('Categorized Questions', 'Organized by topic for easy navigation'),
            ('Typing Indicator', 'Visual feedback while bot generates response'),
            ('Confirmation Modals', 'Prevents accidental clearing of chat history'),
            ('Dark Mode Support', 'Automatic theme based on system preferences'),
            ('Health Monitoring', 'Endpoints for uptime monitoring services')
        ]

        for feat, desc in features:
            p = self.doc.add_paragraph()
            p.add_run(f'{feat}: ').bold = True
            p.add_run(desc)

        self.doc.add_heading('8.2 Chatbot Logic Flow', 3)

        logic_flow = '''
The chatbot uses a rule-based intent matching algorithm:

1. User Input Processing:
   - Convert to lowercase
   - Remove extra whitespace
   - Strip punctuation

2. Intent Matching Priority:
   - Level 1: Greetings (hello, hi, hey) → Return greeting response
   - Level 2: Thanks (thank, thanks) → Return thanks response
   - Level 3: Keyword matching against 25 intents

3. Keyword Matching Algorithm:
   - For each intent, count matching keywords
   - Track intent with highest keyword match count
   - Return response from best matching intent
   - Minimum threshold: at least 1 keyword match

4. Fallback Handling:
   - If no intent matches, return fallback message
   - Fallback includes link to contact support
   - Suggests trying rephrased questions
'''
        self.doc.add_paragraph(logic_flow)

        self.doc.add_heading('8.3 Fallback Logic', 3)
        fallback_text = '''
The fallback system is triggered when:

- No keywords match any intent
- User message is too short or unclear
- Query is outside router support domain

Fallback response: "I'm sorry, I couldn't help with that. Please visit our support page or contact customer service for further assistance. You can also try rephrasing your question or check our suggested questions below."

This ensures users always receive helpful guidance even when the bot cannot answer directly.
'''
        self.doc.add_paragraph(fallback_text)

        self.doc.add_heading('8.4 How to Extend the Project', 3)

        extensions = [
            ('Add New Intents', 'Add new objects to responses array in responses.json with keywords and response'),
            ('Enhance Matching', 'Implement machine learning classifier using scikit-learn or TensorFlow'),
            ('Multi-language Support', 'Add language detection and response localization'),
            ('User Authentication', 'Add login system to save history across devices'),
            ('Admin Dashboard', 'Create interface to manage responses and view analytics'),
            ('WebSocket Support', 'Implement real-time chat using Django Channels'),
            ('Voice Input', 'Add speech recognition for hands-free interaction'),
            ('Image Support', 'Allow users to upload router screenshots for diagnosis')
        ]

        for ext, desc in extensions:
            p = self.doc.add_paragraph()
            p.add_run(f'{ext}: ').bold = True
            p.add_run(desc)

        self.doc.add_heading('8.5 Future Improvements', 3)

        improvements = [
            ('Machine Learning Integration',
             'Replace rule-based matching with NLP model for better intent understanding'),
            ('Sentiment Analysis', 'Detect user frustration and escalate to human support'),
            ('Multi-turn Conversations', 'Maintain context across multiple messages for complex troubleshooting'),
            ('User Feedback Loop', 'Allow users to rate responses to improve accuracy'),
            ('Analytics Dashboard', 'Track common questions, satisfaction rates, and response times'),
            ('WhatsApp Integration', 'Extend chatbot to WhatsApp Business API'),
            ('Voice Assistant', 'Add text-to-speech for accessibility'),
            ('Offline Mode', 'Cache responses for offline availability using Service Workers')
        ]

        for imp in improvements:
            self.add_bullet_point(imp)

    def generate_guide(self):
        """Generate complete Word document guide"""
        print("Generating Django Chatbot Guide...")

        # Create all chapters
        self.create_title_page()
        self.create_table_of_contents()
        self.chapter1_overview()
        self.chapter2_setup()
        self.chapter3_configuration()
        self.chapter4_chatbot_app()
        self.chapter5_templates()
        self.chapter6_static_files()
        self.chapter7_deployment()
        self.chapter8_features()

        # Save the document
        filename = "Complete_Django_Chatbot_Guide.docx"
        self.doc.save(filename)
        print(f"✅ Guide generated successfully: {filename}")
        print(f"📁 File saved in: {os.path.abspath(filename)}")

        return filename


if __name__ == "__main__":
    # Install required package if not present
    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx...")
        os.system("pip install python-docx")
        from docx import Document

    # Generate the guide
    generator = DjangoChatbotGuideGenerator()
    generator.generate_guide()